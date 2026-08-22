"""导出成卷：把题目列表渲染成 Word（.docx），题目卷 + 答案卷分页。

题目/解析里含 ``` 代码块（等宽字体）和 **bold**（Markdown 加粗），都渲染成对应格式。
"""
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor

_CODE_RE = re.compile(r"```\w*\n([\s\S]*?)```")
_BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")


def _add_paragraph(doc: Document, text: str):
    """加段落，把 **bold** 转成 bold run。"""
    p = doc.add_paragraph()
    for part in _BOLD_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def _add_markdown(doc: Document, text: str) -> None:
    """把含 ``` 代码块和 **bold** 的文本加到 docx。"""
    pos = 0
    for m in _CODE_RE.finditer(text):
        normal = text[pos:m.start()].strip()
        if normal:
            for line in normal.split("\n"):
                if line.strip():
                    _add_paragraph(doc, line)
        code = m.group(1)
        for line in code.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(line if line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pos = m.end()
    rest = text[pos:].strip()
    if rest:
        for line in rest.split("\n"):
            if line.strip():
                _add_paragraph(doc, line)


def build_paper_docx(questions: list) -> bytes:
    """把题目列表（Question 对象）渲染成 docx 字节：题目卷 + 答案卷。"""
    doc = Document()
    doc.add_heading("AP Computer Science A 练习卷", 0)

    # 题目卷
    doc.add_heading("题目卷", 1)
    if not questions:
        doc.add_paragraph("（空）")
    for i, q in enumerate(questions, 1):
        doc.add_heading(f"{i}. [{q.type}] {q.difficulty} · {q.knowledge_point}", 2)
        _add_markdown(doc, q.stem)
        if q.type == "选择题" and q.options:
            for j, o in enumerate(q.options):
                _add_paragraph(doc, f"{chr(65 + j)}. {o}")

    # 分页 → 答案卷
    doc.add_page_break()
    doc.add_heading("答案卷", 1)
    for i, q in enumerate(questions, 1):
        p = _add_paragraph(doc, f"{i}. 答案：{q.answer}")
        for run in p.runs:
            run.bold = True
        if q.explanation:
            _add_paragraph(doc, "解析：")
            _add_markdown(doc, q.explanation)
        if q.marking_points:
            _add_paragraph(doc, "评分要点：")
            for mp in q.marking_points:
                _add_paragraph(doc, f"• {mp}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
